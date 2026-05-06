"""
生成 Word 周报：格式与模板完全一致，内容来自 entries.json
用法: python scripts/generate_report.py
"""

import json
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE_DIR = os.path.dirname(os.path.dirname(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'data', 'entries.json')

# ========== 样式常量（与模板完全一致）==========
HEADING1_SIZE = Pt(18)
HEADING1_LINE_SPACING = Pt(36)
HEADING2_SIZE = Pt(15)
HEADING2_LINE_SPACING = Pt(30)
HEADING3_SIZE = Pt(14)
HEADING3_LINE_SPACING = Pt(28)
NORMAL_SIZE = Pt(10.5)
NORMAL_LINE_SPACING = Pt(21)
PAGE_WIDTH = 7556500            # EMU ≈ A4
PAGE_HEIGHT = 10693400          # EMU ≈ A4
MARGIN = 914400                 # EMU = 1 inch

FONT_WESTERN = 'Times New Roman'
FONT_EASTERN = '宋体'

STAR_MARKER = '★'  # 重要标记


def set_font(run, size=NORMAL_SIZE, bold=None, color=None, western=FONT_WESTERN, eastern=FONT_EASTERN):
    """设置 run 的字体（中西文分别指定）"""
    run.font.size = size
    run.font.name = western
    # 设置中文字体
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = r.makeelement(qn('w:rPr'), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), eastern)
    rFonts.set(qn('w:ascii'), western)
    rFonts.set(qn('w:hAnsi'), western)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(para, line_spacing, space_before=0, space_after=0):
    """设置段落行距和段前段后"""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_title(doc, text):
    """添加 Heading 1 标题"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, HEADING1_LINE_SPACING)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font(run, size=HEADING1_SIZE, bold=True)


def add_section_title(doc, text):
    """添加 Heading 2 节标题"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, HEADING2_LINE_SPACING, space_before=8)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_font(run, size=HEADING2_SIZE, bold=True)


def add_entry_title(doc, index, country, title, important=False):
    """添加条目标题: 【序号】国家 | 标题[★]"""
    star = STAR_MARKER if important else ''
    text = f'【{index}】{country} | {title}{star}'
    para = doc.add_paragraph()
    set_paragraph_spacing(para, HEADING3_LINE_SPACING, space_before=6)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_font(run, size=HEADING3_SIZE, bold=True)


def add_content(doc, text):
    """添加正文段落（两端对齐）"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, NORMAL_LINE_SPACING)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # 首行缩进2字符
    para.paragraph_format.first_line_indent = Pt(21)
    run = para.add_run(text)
    set_font(run, size=NORMAL_SIZE)
    return para


def add_source(doc, text):
    """添加资料来源行"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, NORMAL_LINE_SPACING, space_before=2)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(f'资料来源：{text}')
    set_font(run, size=NORMAL_SIZE)
    return para


def get_report_period(entries):
    """从条目日期推断报告周期"""
    dates = sorted(e['date'] for e in entries)
    if not dates:
        today = date.today()
        return today.strftime('%Y%m%d'), today.strftime('%Y%m%d')
    return dates[0].replace('-', ''), dates[-1].replace('-', '')


def generate_report():
    # 加载数据
    with open(DATA_FILE, 'r', encoding='utf-8') as f:
        data = json.load(f)
    entries = data['entries']

    start_date, end_date = get_report_period(entries)
    title = f'养老金资讯-国际养老金动态（{start_date}-{end_date}）'

    # 按分类分组
    categories = {}
    for e in entries:
        categories.setdefault(e['category'], []).append(e)

    # 定义章节顺序
    section_order = [
        ('政策改革', '一、养老金政策改革'),
        ('政策调整', '二、养老金政策调整'),
        ('投资动态', '三、养老金基金投资'),
        ('数据发布', '四、数据发布与统计'),
        ('法规更新', '五、法规更新'),
        ('行政安排', '六、行政安排'),
    ]

    # 创建文档
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

    # === 标题 ===
    add_title(doc, title)

    # 空行
    p1 = doc.add_paragraph()

    # ================================================================
    #  目  录
    # ================================================================
    toc_heading = doc.add_paragraph()
    set_paragraph_spacing(toc_heading, HEADING2_LINE_SPACING)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_heading.add_run('目  录')
    set_font(run, size=HEADING2_SIZE, bold=True)

    toc_entries_para = doc.add_paragraph()
    set_paragraph_spacing(toc_entries_para, NORMAL_LINE_SPACING, space_before=4)
    toc_entries_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    entry_counter = 0
    toc_lines = []
    for cat_key, section_title in section_order:
        cat_entries = categories.get(cat_key, [])
        if not cat_entries:
            continue
        cat_entries.sort(key=lambda e: e['date'], reverse=True)

        line = f'{section_title}（共 {len(cat_entries)} 条）'
        toc_lines.append(line)

        for e in cat_entries:
            entry_counter += 1
            important = e.get('importance', 3) >= 4
            star = STAR_MARKER if important else ''
            toc_lines.append(f'    【{entry_counter}】{e["country"]} | {e["title"]}{star}')

    # 将目录项合并为一个段落（用换行符分隔）
    run = toc_entries_para.add_run('\n'.join(toc_lines))
    set_font(run, size=NORMAL_SIZE)

    # 空行分隔
    p2 = doc.add_paragraph()

    # ================================================================
    #  正文内容
    # ================================================================

    # 重置条目编号
    entry_counter = 0
    for cat_key, section_title in section_order:
        cat_entries = categories.get(cat_key, [])
        if not cat_entries:
            continue

        # 按日期排序（最新的在前）
        cat_entries.sort(key=lambda e: e['date'], reverse=True)

        add_section_title(doc, section_title)

        for e in cat_entries:
            entry_counter += 1
            important = e.get('importance', 3) >= 4
            add_entry_title(doc, entry_counter, e['country'], e['title'], important)
            add_content(doc, e['content'])
            add_source(doc, f"{e['source']}（{e['date']}）")

    # 保存
    output_dir = os.path.join(BASE_DIR, 'data', 'reports')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'国际养老金动态（{start_date}-{end_date}）.docx')
    doc.save(output_path)
    print(f'报告已生成：{output_path}')
    print(f'共 {entry_counter} 条动态，覆盖 {len(categories)} 个分类')


if __name__ == '__main__':
    generate_report()
